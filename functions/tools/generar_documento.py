"""Tool generar_documento: compone la historia en el formato de plantilla.

Genera siempre Markdown y, si python-docx está instalado, también un DOCX
con las seis tablas de dos filas que usan las historias de ejemplo.
Los borradores se escriben en `output/` y la tool devuelve sus rutas.
"""

import re
import unicodedata
from pathlib import Path

OUTPUT_DIR = Path("/tmp/hu-output")


def _sanitizar_nombre(texto: str) -> str:
    nfkd = unicodedata.normalize("NFKD", texto)
    sin_acentos = "".join(c for c in nfkd if not unicodedata.combining(c))
    return re.sub(r"[^\w\-]", "_", sin_acentos).strip("_")

ETIQUETAS = [
    ("usuario", "Usuario ¿Quién?", "Usuarios de negocio que participan en la historia"),
    ("funcionalidad", "Funcionalidad ¿Qué?", "Descripción corta de la historia de usuario."),
    ("objetivo", "Objetivo ¿Para qué?", "Para qué se necesita la historia de usuario."),
    ("descripcion", "Descripción ¿Cómo?", "Descripción de cómo se debe llevar a cabo la historia de usuario"),
    ("criterios_aceptacion", "Criterios de aceptación", "Condiciones específicas que deben cumplirse una vez desarrollada"),
    ("aceptacion_usuarios", "Aceptación Usuarios", "Especificar Quién y Cuándo"),
]


def _lineas_descripcion(descripcion: dict) -> list[str]:
    lineas: list[str] = []
    if descripcion.get("encuadre"):
        lineas.append(descripcion["encuadre"])
    lineas.extend(descripcion.get("precondiciones", []))
    lineas.extend(f"{i}. {paso}" for i, paso in enumerate(descripcion.get("pasos", []), start=1))
    lineas.extend(descripcion.get("validaciones", []))
    if descripcion.get("almacenamiento"):
        lineas.append(descripcion["almacenamiento"])
    lineas.extend(descripcion.get("casos_error", []))
    return lineas


def _lineas_bloque(clave: str, historia: dict) -> list[str]:
    valor = historia.get(clave, "")
    if clave == "descripcion":
        return _lineas_descripcion(valor or {})
    if clave == "criterios_aceptacion":
        return [str(c) for c in (valor or [])]
    if clave == "aceptacion_usuarios":
        valor = valor or {}
        return [f"Quién: {valor.get('quien', '')}", f"Cuándo: {valor.get('cuando', '')}"]
    return [str(valor)]


def _render_markdown(historia: dict) -> str:
    titulo = f"{historia.get('identificador', '')} - {historia.get('titulo', '')}".strip(" -")
    partes = [f"# {titulo}", "", "## Historia de Usuario", ""]
    for clave, etiqueta, guia in ETIQUETAS:
        partes.append(f"| {etiqueta} | {guia} |")
        partes.append("|---|---|")
        contenido = "<br>".join(_lineas_bloque(clave, historia))
        partes.append(f"| | {contenido} |")
        partes.append("")
    return "\n".join(partes)


def _render_docx(historia: dict, ruta: Path) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return False

    doc = Document()
    titulo = f"{historia.get('identificador', '')} - {historia.get('titulo', '')}".strip(" -")
    doc.add_heading(titulo, level=1)
    doc.add_heading("Historia de Usuario", level=2)

    for clave, etiqueta, guia in ETIQUETAS:
        tabla = doc.add_table(rows=2, cols=2)
        tabla.style = "Table Grid"
        celda_etiqueta = tabla.cell(0, 0)
        celda_etiqueta.text = etiqueta
        for run in celda_etiqueta.paragraphs[0].runs:
            run.bold = True
        tabla.cell(0, 1).text = guia
        tabla.cell(1, 0).text = ""
        celda_contenido = tabla.cell(1, 1)
        lineas = _lineas_bloque(clave, historia)
        celda_contenido.text = lineas[0] if lineas else ""
        for linea in lineas[1:]:
            parrafo = celda_contenido.add_paragraph(linea)
            parrafo.paragraph_format.space_after = Pt(0)
        doc.add_paragraph("")

    doc.save(str(ruta))
    return True


def generar_documento(historia: dict) -> dict:
    OUTPUT_DIR.mkdir(exist_ok=True)
    id_raw = historia.get("identificador", "HU-XXX")
    titulo_raw = historia.get("titulo", "sin-titulo")
    base = _sanitizar_nombre(f"{id_raw}-{titulo_raw}")

    ruta_md = OUTPUT_DIR / f"{base}.md"
    ruta_md.write_text(_render_markdown(historia), encoding="utf-8")

    resultado = {"markdown": str(ruta_md), "docx": None}
    ruta_docx = OUTPUT_DIR / f"{base}.docx"
    if _render_docx(historia, ruta_docx):
        resultado["docx"] = str(ruta_docx)
    else:
        resultado["aviso"] = "python-docx no está instalado; solo se ha generado Markdown."
    return resultado
